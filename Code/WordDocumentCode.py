# import packages (make sure python-docx is installed)
import docx

# Open the input Word document
input_doc = docx.Document("/content/ADHD Test document.docx") #Add your file path here

# Create a new Word document to store the amended words
output_doc = docx.Document()

# Loop through each paragraph in the input document
for paragraph in input_doc.paragraphs:

    # Create a new paragraph in the output document
    output_paragraph = output_doc.add_paragraph()

    # Loop through each run in the input paragraph
    for run in paragraph.runs:

        # Split the run text into individual words
        words = run.text.split()

        # Loop through each word and bold the first half
        for word in words:
            half_len = len(word) // 2
            output_run = output_paragraph.add_run()
            output_run.text = word[:half_len]
            output_run.bold = True
            output_run = output_paragraph.add_run()
            output_run.text = word[half_len:] + ' '

        # Add a space character after the run
        output_paragraph.add_run(' ')

    # Add a new line after the paragraph
    output_doc.add_paragraph()

# Name output file
output_name = input('Please input name of output file: ')
# Save the output document
output_doc.save(output_name+".docx")