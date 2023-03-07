#WORK IN PROGRESS!!! NEVER BEEN TESTED
import io
import os

from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload, MediaFileUpload

import docx

# Set up the Google Drive API client
SCOPES = ['https://www.googleapis.com/auth/drive']
SERVICE_ACCOUNT_FILE = 'service_account.json'
creds = None
creds = service_account.Credentials.from_service_account_file(
        SERVICE_ACCOUNT_FILE, scopes=SCOPES)
service = build('drive', 'v3', credentials=creds)

# ID of the Google Doc you want to modify
doc_id = 'YOUR_GOOGLE_DOC_ID_HERE'

# Download the Google Doc as a Word document
file_id = service.files().export(fileId=doc_id, mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document').execute().get('id')
request = service.files().get_media(fileId=file_id)
fh = io.BytesIO()
downloader = MediaIoBaseDownload(fh, request)
done = False
while done is False:
    status, done = downloader.next_chunk()
    print(f'Download {int(status.progress() * 100)}.')
fh.seek(0)

# Open the Word document
input_doc = docx.Document(fh)

# Create a new Word document to store the amended words
output_doc = docx.Document()

# Loop through each paragraph in the input document
for paragraph in input_doc.paragraphs:

    # Create a new paragraph in the output document
    output_paragraph = output_doc.add_paragraph()

    # Loop through each run in the input paragraph
    for run in paragraph.runs:

        # Split the run text into individual words
        words = run.text.split()

        # Loop through each word and bold the first half
        for word in words:
            half_len = len(word) // 2
            output_run = output_paragraph.add_run()
            output_run.text = word[:half_len]
            output_run.bold = True
            output_run = output_paragraph.add_run()
            output_run.text = word[half_len:] + ' '

        # Add a space character after the run
        output_paragraph.add_run(' ')

    # Add a new line after the paragraph
    output_doc.add_paragraph()

# Save the output document to a file
output_doc.save('output.docx')

# Upload the modified Word document back to Google Drive
file_metadata = {'name': 'output.docx', 'parents': [doc_id]}
media = MediaFileUpload('output.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessing')
